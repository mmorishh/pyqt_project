import sqlite3
import sys

from docx import Document
from docx.shared import Inches
from PyQt5.QtGui import QPixmap
from PyQt5.QtWidgets import QMainWindow, QApplication
from PyQt5.QtCore import Qt

from genshin_pyqt_ui import Ui_MainWindow


class MyWidget(QMainWindow, Ui_MainWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.con = sqlite3.connect("genshin_db.sqlite")
        cur = self.con.cursor()

        self.cb_elements.addItems(
            [item[0] for item in cur.execute("SELECT name FROM elements").fetchall()])
        self.cb_weapons.addItems(
            [item[0] for item in cur.execute("SELECT name FROM weap_role").fetchall()])
        self.btn_show.clicked.connect(self.filter)

        self.btn_download.clicked.connect(self.download)

        self.name = 'Diluc'

        res = cur.execute("SELECT artifacts FROM characters WHERE name LIKE ?", (self.name,)).fetchall()
        for i in res:
            for j in i:
                res1 = j.split(' & ')
                for i in res1:
                    self.listWidget.addItem('Weapon: ' + i)

    def search(self):  #  работает с поиском
        pass


    def filter(self):
        self.listWidget_2.clear()
        cur = self.con.cursor()
        res_element = cur.execute(f"""
                      SELECT characters.name
                      FROM characters, elements JOIN characters_elements
                      ON characters.id = characters_elements.id_elements
                      WHERE id_elements = {self.cb_elements.currentIndex() + 1}
                      """).fetchall()
        for i in res_element:
            for j in i:
                for _ in j:
                    self.listWidget_2.addItem(_)


        res_weapon = cur.execute(f"""
                      SELECT characters.id
                      FROM characters, weap_role JOIN weap_role_characters
                      ON characters.id = weap_role_characters.id_weap_role
                      WHERE id_weap_role = {self.cb_weapons.currentIndex() + 1}
                      """).fetchall()

    def btn(self):
        self.con = sqlite3.connect("genshin_db.sqlite")
        cur = self.con.cursor()

        res = cur.execute("SELECT name FROM characters WHERE id=1").fetchall()
        for i in res:
            for j in i:
                self.listWidget.addItem('Name: ' + j)

        res1 = cur.execute("SELECT element FROM characters WHERE id=1").fetchall()
        for i in res1:
            for j in i:
                self.listWidget.addItem('Element: ' + j)

        res2 = cur.execute("SELECT weap_role FROM characters WHERE id=1").fetchall()
        for i in res2:
            for j in i:
                self.listWidget.addItem('Weapon: ' + j)

        res3 = cur.execute("SELECT artifacts FROM characters WHERE id=1").fetchall()
        for i in res3:
            for j in i:
                self.listWidget.addItem('Artifacts: ' + j)

        res4 = cur.execute("SELECT weapons FROM characters WHERE id=1").fetchall()
        for i in res4:
            for j in i:
                self.listWidget.addItem('Weapons: ' + j)

        res5 = cur.execute("SELECT image FROM characters WHERE id=1").fetchall()
        for i in res5:
            for j in i:
                pixmap = QPixmap(j)
                self.label_image.setPixmap(pixmap)

    def download(self):
        if self.name != '':
            self.con = sqlite3.connect("genshin_db.sqlite")
            cur = self.con.cursor()

            document = Document()

            document.add_heading(self.name, 0)

            document.add_heading('Element', level=1)
            res = cur.execute("SELECT element FROM characters WHERE name LIKE ?", (self.name,)).fetchall()
            for i in res:
                for j in i:
                    p = document.add_paragraph(j)

            document.add_heading('Weapon', level=1)
            res = cur.execute("SELECT weap_role FROM characters WHERE name LIKE ?", (self.name,)).fetchall()
            for i in res:
                for j in i:
                    p = document.add_paragraph(j)

            document.add_heading('Artifacts', level=1)
            res = cur.execute("SELECT artifacts FROM characters WHERE name LIKE ?", (self.name,)).fetchall()
            for i in res:
                for j in i:
                    res1 = j.split(' & ')
                    for i in res1:
                        document.add_paragraph(i,
                                               style='List Bullet')

            document.add_heading('Weapons', level=1)
            res = cur.execute("SELECT weapons FROM characters WHERE name LIKE ?", (self.name,)).fetchall()
            for i in res:
                for j in i:
                    res1 = j.split(' or ')
                    for i in res1:
                        document.add_paragraph(i,
                                               style='List Bullet')

            res = cur.execute("SELECT image FROM characters WHERE name LIKE ?", (self.name,)).fetchall()
            for i in res:
                for j in i:
                    document.add_picture(j)

            document.save(self.name + '.docx')


def main():
    app = QApplication(sys.argv)
    ex = MyWidget()
    ex.show()
    sys.exit(app.exec_())


if __name__ == '__main__':
    main()
